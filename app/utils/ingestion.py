"""Extract plain text from uploads (PDF / Word / plain text)."""

from pathlib import Path


def extract_text_from_file(path: Path, mime_type: str | None) -> str:
    """Use file suffix when browsers send generic MIME types (e.g. octet-stream)."""
    suf = path.suffix.lower()
    mt = (mime_type or "").lower()

    if suf == ".pdf" or "pdf" in mt:
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(str(path))
            parts: list[str] = []
            for page in reader.pages:
                parts.append(page.extract_text() or "")
            return "\n".join(parts)
        except Exception:
            return ""

    if suf == ".docx" or "wordprocessingml" in mt or "officedocument.wordprocessingml" in mt:
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(str(path))
            parts: list[str] = []
            for p in doc.paragraphs:
                t = (p.text or "").strip()
                if t:
                    parts.append(t)
            for table in doc.tables:
                for row in table.rows:
                    cells = [(c.text or "").strip() for c in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except Exception:
            return ""

    raw = path.read_bytes()
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("utf-8", errors="ignore")


# Text budget per imaging file when building chat context (MRI reports as PDF, etc.)
_MAX_IMAGING_EXCERPT_CHARS = 6000


def build_imaging_context_excerpt(
    path: Path,
    *,
    filename: str,
    mime_type: str | None,
    caption: str | None,
    image_id: int,
) -> str:
    """Single study/file block for the clinical assistant (PDF text + metadata)."""
    lines: list[str] = [
        f"[Imaging upload id={image_id}]",
        f"File name: {filename}",
    ]
    if caption and str(caption).strip():
        lines.append(f"Staff caption / clinical note: {str(caption).strip()}")

    suf = path.suffix.lower()
    mt = (mime_type or "").lower()

    if suf == ".pdf" or "pdf" in mt:
        extracted = extract_text_from_file(path, mime_type)
        if extracted.strip():
            t = extracted.strip()
            if len(t) > _MAX_IMAGING_EXCERPT_CHARS:
                t = t[:_MAX_IMAGING_EXCERPT_CHARS] + "\n… [truncated]"
            lines.append("Extracted text from imaging PDF (reports, typed summaries):")
            lines.append(t)
        else:
            lines.append(
                "(PDF on file; no extractable text — it may be image-only scans. Add a caption or upload a text report.)"
            )
    elif suf in (".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tiff", ".tif") or any(
        x in mt for x in ("image/png", "image/jpeg", "image/webp", "image/gif", "image/bmp", "image/tiff")
    ):
        lines.append(
            "(Raster image on file — describe key findings in the caption, or upload a PDF/text radiology report for full-text retrieval.)"
        )
    elif suf == ".dcm" or "dicom" in mt:
        lines.append("(DICOM on file; binary pixel data is not decoded here — use caption or derived PDF report.)")
    else:
        lines.append(f"(File on file; MIME {mime_type or 'unknown'} — no automated text extraction.)")

    return "\n".join(lines)
